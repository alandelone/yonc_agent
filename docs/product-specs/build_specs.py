from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"

SOURCE_CHAT = "https://chatgpt.com/share/6a90f2cc-41b0-83ec-b4e9-19fac42969e4"
SOURCE_REPO = "https://github.com/alandelone/yonc_agent"
SOURCE_UI_REFERENCE = "https://dribbble.com/shots/25813944-FlowTune-Dashboard-Optimizing-Flow-Processes"
SOURCE_MD = (ROOT.parents[1] / "Yonc_Graph_Project_System_Spec_v0.1.md").resolve().as_uri()

NAVY = "17324D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "17212B"
MUTED = "667482"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
CORAL = "C65F58"
CYAN = "3A91B5"
GOLD = "9B761C"
RED = "9B1C1C"
GREEN = "356859"

PAGE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, name: str = "Calibri", size: float | None = None,
                 color: str | None = None, bold: bool | None = None,
                 italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 6,
                          line: float = 1.25, keep_with_next: bool | None = None) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if keep_with_next is not None:
        paragraph.paragraph_format.keep_with_next = keep_with_next


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80,
                     start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "D7DEE6", size: int = 6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != PAGE_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {PAGE_WIDTH_DXA}, got {sum(widths_dxa)}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row_idx, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if row_idx == 0:
            repeat_header = OxmlElement("w:tblHeader")
            repeat_header.set(qn("w:val"), "true")
            tr_pr.append(repeat_header)
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_hyperlink(paragraph, text: str, url: str, color: str = BLUE) -> None:
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_color = OxmlElement("w:color")
    r_color.set(qn("w:val"), color)
    r_pr.append(r_color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(fonts)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_numbering_definition(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


class SpecDoc:
    def __init__(self, title: str, subtitle: str, short_title: str, accent: str = BLUE):
        self.doc = Document()
        self.title = title
        self.subtitle = subtitle
        self.short_title = short_title
        self.accent = accent
        self.bullet_num_id = add_numbering_definition(self.doc, "bullet")
        self.decimal_num_id = add_numbering_definition(self.doc, "decimal")
        self._configure_document()

    def _configure_document(self) -> None:
        doc = self.doc
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        styles = doc.styles
        normal = styles["Normal"]
        normal.font.name = "Calibri"
        normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        normal.font.size = Pt(11)
        normal.font.color.rgb = rgb(INK)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25

        title_style = styles["Title"]
        title_style.font.name = "Calibri"
        title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        title_style.font.size = Pt(30)
        title_style.font.bold = True
        title_style.font.color.rgb = rgb(NAVY)
        title_style.paragraph_format.space_before = Pt(0)
        title_style.paragraph_format.space_after = Pt(8)

        subtitle_style = styles["Subtitle"]
        subtitle_style.font.name = "Calibri"
        subtitle_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        subtitle_style.font.size = Pt(14)
        subtitle_style.font.color.rgb = rgb(MUTED)
        subtitle_style.paragraph_format.space_before = Pt(0)
        subtitle_style.paragraph_format.space_after = Pt(18)

        heading_specs = {
            "Heading 1": (16, BLUE, 18, 10),
            "Heading 2": (13, BLUE, 14, 7),
            "Heading 3": (12, DARK_BLUE, 10, 5),
        }
        for style_name, (size, color, before, after) in heading_specs.items():
            style = styles[style_name]
            style.font.name = "Calibri"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = rgb(color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True

        caption = styles["Caption"]
        caption.font.name = "Calibri"
        caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        caption.font.size = Pt(9)
        caption.font.color.rgb = rgb(MUTED)
        caption.font.italic = True
        caption.paragraph_format.space_before = Pt(4)
        caption.paragraph_format.space_after = Pt(10)
        caption.paragraph_format.keep_with_next = False

        self._set_header_footer(section)
        doc.core_properties.title = self.title
        doc.core_properties.subject = self.subtitle
        doc.core_properties.author = "Yonc Project Documentation"

    def _set_header_footer(self, section) -> None:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(self.short_title)
        set_run_font(r, size=8.5, color=MUTED, bold=True)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after = Pt(0)
        fr = fp.add_run("Yonc Product Specification  |  ")
        set_run_font(fr, size=8.5, color=MUTED)
        add_field(fp, "PAGE")

    def add_cover(self, edition: str, audience: str) -> None:
        for _ in range(4):
            p = self.doc.add_paragraph()
            set_paragraph_spacing(p, after=10)
        kicker = self.doc.add_paragraph()
        kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(kicker, after=16)
        kr = kicker.add_run("YONC PROJECT SYSTEM")
        set_run_font(kr, size=10, color=self.accent, bold=True)

        title = self.doc.add_paragraph(style="Title")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run(self.title)

        subtitle = self.doc.add_paragraph(style="Subtitle")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.add_run(self.subtitle)

        meta = self.doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(meta, after=6)
        mr = meta.add_run(f"{edition}  |  Audience: {audience}")
        set_run_font(mr, size=10, color=MUTED)

        purpose = self.doc.add_paragraph()
        purpose.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(purpose, before=28, after=8)
        pr = purpose.add_run("Lossless structured synthesis of the shared design conversation")
        set_run_font(pr, size=11, color=INK, italic=True)

        source = self.doc.add_paragraph()
        source.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(source, after=0)
        add_hyperlink(source, "Shared design conversation", SOURCE_CHAT)
        self.doc.add_page_break()

    def h1(self, text: str) -> None:
        self.doc.add_paragraph(text, style="Heading 1")

    def h2(self, text: str) -> None:
        self.doc.add_paragraph(text, style="Heading 2")

    def h3(self, text: str) -> None:
        self.doc.add_paragraph(text, style="Heading 3")

    def p(self, text: str = "", bold_lead: str | None = None) -> None:
        p = self.doc.add_paragraph()
        set_paragraph_spacing(p)
        if bold_lead and text.startswith(bold_lead):
            r1 = p.add_run(bold_lead)
            set_run_font(r1, bold=True)
            r2 = p.add_run(text[len(bold_lead):])
            set_run_font(r2)
        else:
            r = p.add_run(text)
            set_run_font(r)

    def bullets(self, items: Iterable[str]) -> None:
        for item in items:
            p = self.doc.add_paragraph()
            apply_numbering(p, self.bullet_num_id)
            set_paragraph_spacing(p, after=4, line=1.25)
            r = p.add_run(item)
            set_run_font(r)

    def numbered(self, items: Iterable[str]) -> None:
        for item in items:
            p = self.doc.add_paragraph()
            apply_numbering(p, self.decimal_num_id)
            set_paragraph_spacing(p, after=4, line=1.25)
            r = p.add_run(item)
            set_run_font(r)

    def code(self, text: str) -> None:
        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before=4, after=8, line=1.0)
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.right_indent = Inches(0.18)
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), LIGHT_GRAY)
        p_pr.append(shd)
        r = p.add_run(text)
        set_run_font(r, name="Consolas", size=9.2, color=INK)

    def callout(self, label: str, text: str, fill: str = CALLOUT) -> None:
        table = self.doc.add_table(rows=1, cols=1)
        set_table_geometry(table, [PAGE_WIDTH_DXA])
        set_table_borders(table, color="D9E1E8", size=5)
        cell = table.cell(0, 0)
        shade_cell(cell, fill)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=0)
        lr = p.add_run(f"{label}: ")
        set_run_font(lr, bold=True, color=NAVY)
        tr = p.add_run(text)
        set_run_font(tr)
        after = self.doc.add_paragraph()
        set_paragraph_spacing(after, after=4)

    def table(self, headers: Sequence[str], rows: Sequence[Sequence[str]], widths_dxa: Sequence[int],
              header_fill: str = LIGHT_BLUE, font_size: float = 9.5) -> None:
        table = self.doc.add_table(rows=1, cols=len(headers))
        set_table_geometry(table, widths_dxa)
        set_table_borders(table)
        hdr = table.rows[0]
        set_repeat_table_header(hdr)
        for i, text in enumerate(headers):
            cell = hdr.cells[i]
            shade_cell(cell, header_fill)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.15)
            r = p.add_run(text)
            set_run_font(r, size=font_size, bold=True, color=NAVY)
        for row_data in rows:
            row = table.add_row()
            for i, value in enumerate(row_data):
                cell = row.cells[i]
                p = cell.paragraphs[0]
                set_paragraph_spacing(p, after=0, line=1.15)
                r = p.add_run(str(value))
                set_run_font(r, size=font_size)
        set_table_geometry(table, widths_dxa)
        after = self.doc.add_paragraph()
        set_paragraph_spacing(after, after=4)

    def image(self, path: Path, caption: str, alt_text: str, width: float = 6.5) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        run = p.add_run()
        shape = run.add_picture(str(path), width=Inches(width))
        inline = shape._inline
        doc_pr = inline.docPr
        doc_pr.set("descr", alt_text)
        cap = self.doc.add_paragraph(caption, style="Caption")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def sources(self, extra: Sequence[tuple[str, str]] = ()) -> None:
        self.h1("Source references")
        entries = [
            ("Shared ChatGPT design conversation", SOURCE_CHAT),
            ("Yonc Graph Project System Specification v0.1", SOURCE_MD),
            ("Yonc Agent repository", SOURCE_REPO),
        ] + list(extra)
        for label, url in entries:
            p = self.doc.add_paragraph()
            apply_numbering(p, self.bullet_num_id)
            set_paragraph_spacing(p, after=4)
            add_hyperlink(p, label, url)

    def save(self, filename: str) -> Path:
        path = ROOT / filename
        self.doc.save(path)
        return path


def build_ui_spec() -> Path:
    s = SpecDoc(
        "Canvas UI & Elastic Timeline View",
        "Interaction, visual grammar, scheduling behavior, and cross-view requirements",
        "Canvas & Timeline UI",
        accent=CORAL,
    )
    s.add_cover("Specification v1.0 — 28 Aug 2026", "Product, UX, frontend, graph/runtime")

    s.h1("Information preservation note")
    s.p("This document is a structured, lossless synthesis of the shared design conversation and the supplied Yonc_Graph_Project_System_Spec_v0.1.md reference. Conversational repetition and speech disfluencies have been normalized, but accepted requirements, examples, corrections, rejected interpretations, deferred ideas, and unresolved mappings are retained and labeled. The dedicated Timeline View is documented separately from the Canvas even though both project the same graph.")
    s.callout("Core product principle", "Show the whole system when the user wants to understand; hide the system when the user wants to act.")

    s.h1("1. View architecture")
    s.p("The system is not one screen. It is one project graph projected into several views, each answering a different question and preserving its own view state.")
    s.table(
        ["View", "Primary question", "Primary unit", "Interaction model"],
        [
            ["Portfolio / Project", "What projects or modules are currently worth managing?", "WBS L1–L2 or any selected project scope", "Cards, filters, scope selection"],
            ["Canvas", "How is the system structured and connected?", "Project through Action, with selective Resources", "Zoom, pan, toggle, constrained drag, split"],
            ["Timeline", "How much time does the work occupy, overlap, and pressure the deadline?", "Mostly WBS L2–L3; actions support estimates", "Drag, resize, day-block scheduling, forecast"],
            ["Mobile / Focus", "What matters now on a small screen?", "Important projects, current actions, deadlines", "Tap, long-press, review; no graph dragging"],
        ],
        [1700, 3300, 2360, 2000],
    )
    s.bullets([
        "Desktop/tablet uses the graph dashboard and full Timeline. Mobile does not shrink the entire Canvas; it exposes focused project/action cards.",
        "Notion remains a possible projection, not the source of truth.",
        "A node may be promoted to the current Project Scope. Project is therefore partly a viewing scale, not only a fixed node type.",
        "WBS L2 modules may appear as project cards in Portfolio View even when their graph node type remains Deliverable or Module.",
        "Canvas and Timeline should normally remember expanded/collapsed state independently because the user may want different detail in each projection.",
    ])

    s.h1("2. Canvas View — spatial graph")
    s.image(
        ASSETS / "canvas-ui-authoritative-reference.png",
        "Figure 1. User-supplied Canvas UI reference: time-constrained graph, orthogonal dependencies, WBS cards, progress, resources, and inspector.",
        "Dark desktop Yonc Canvas reference showing a temporal graph axis, orthogonal edges, project and WBS nodes, lifecycle treatments, resources, progress, deadlines, temporal pressure, and a node inspector.",
    )
    s.h2("2.1 Canvas purpose and shell")
    s.p("The Canvas is a structured, time-constrained graph surface, not a generic whiteboard and not a text editor. It supports understanding, planning, structural navigation, and graph construction. Detailed editing belongs in an inspector or linear detail surface.")
    s.table(
        ["Region", "Required content"],
        [
            ["Top toolbar", "Project/scope selector, time context, filter, search, settings; quarter markers may appear as quiet reference points."],
            ["Left rail", "Portfolio, Canvas, Timeline, Focus, Search, Auto-layout, Zoom."],
            ["Central Canvas", "Project containers, nodes, orthogonal edges, temporal constraints, selected neighborhood."],
            ["Lower-left minimap", "Current location inside a large project landscape; retain because the graph may span many scopes."],
            ["Right inspector", "Selected node facts, status/stage, time fields, dependencies, resources, warnings, and events."],
            ["Bottom strip", "Restrained metrics only: progress, blocked count, deadline/temporal pressure, recent activity if useful."],
        ],
        [2100, 7260],
    )

    s.h2("2.2 Spatial and temporal rules")
    s.bullets([
        "The horizontal axis carries time: NOW on the left and future/deadline toward the right. The vertical axis separates project branches and hierarchy.",
        "The background uses very faint quarterly reference lines (for example Mar, Jun, Sep, Dec). Months or weeks appear only when zoomed in; never render an Excel-like full grid on the Canvas.",
        "A node with an explicit deadline is an Anchor Node. Its horizontal position is locked on the Canvas; changing it requires the node detail flow, not casual dragging.",
        "A child without its own deadline is a Flexible Node. It may move inside its ancestor's allowed temporal region but cannot cross the nearest deadline-bearing ancestor.",
        "Parent deadline is the descendants' right boundary. Dependency constraints also apply: if B depends on A, B cannot be placed before A. Forbidden regions appear during drag.",
        "Dragging changes planned/rough temporal placement and view position; it never infers or changes parent-child hierarchy.",
        "Start with auto-layout. Manual movement refines a stable structure; free dragging must not become the main organization mechanism.",
    ])
    s.code("Structural: Project → Deliverable → Work Package → Action\nLogical:   A depends_on B\nTemporal:  child.end ≤ parent.deadline; B.start ≥ A.complete")

    s.h2("2.3 Node grammar")
    s.p("Each visual channel owns one meaning. This avoids the common failure where red could mean project, blocked, urgent, or overdue at the same time.")
    s.table(
        ["Visual channel", "Meaning", "Rule"],
        [
            ["Project hue", "Which project/line owns the node", "One hue family per project; do not reuse hue as status."],
            ["Color lightness", "WBS/decomposition depth", "Project/Goal darkest; Action lightest."],
            ["Whole-card border", "Lifecycle/stage/runtime state", "Use perimeter treatment, not a small status dot."],
            ["Glow intensity", "Temporal pressure / urgency", "Derived from Timeliner; glow is information, never decoration."],
            ["Warning icon", "Graph health issue", "Small corner signal; details appear on click/hover."],
            ["Expand/collapse", "How much structure the user currently wants", "Persistent view state, separate from graph facts."],
            ["Canvas position", "Rough temporal planning and branch location", "Constrained by deadline and dependency rules."],
        ],
        [2200, 2700, 4460],
    )

    s.h2("2.4 Lifecycle border and light language")
    s.table(
        ["State", "Block treatment", "Motion / attention"],
        [
            ["Planning", "Dim body; thin or partially illuminated perimeter", "Static and quiet"],
            ["Proposed / discussion", "Semi-transparent body; incomplete glow; dashed proposal edges", "Subtle only; visually distinct from committed reality"],
            ["Approved / Ready", "Complete clean border", "Bright enough to read, normally static"],
            ["Active / Doing", "Complete border and full-card aura", "Slow breathing glow; do not use gaming-style RGB animation"],
            ["Blocked", "Broken or interrupted border/glow", "Static; flow interruption is visible without turning the whole card red"],
            ["Review", "Complete border with a thin perimeter sweep", "Restrained slow motion"],
            ["Done", "Low-contrast body and no glow", "Static; completed work should stop competing for attention"],
        ],
        [1500, 5160, 2700],
    )
    s.p("The source discussion briefly explored a small punk/cyber emblem (black hair, one eye, crown) in a card corner. That idea was not accepted as the status mechanism; the user corrected the design toward whole-card perimeter light. Treat an emblem/character mark as optional future identity styling, never as the primary lifecycle signal.")

    s.h2("2.5 Attention and motion budget")
    s.bullets([
        "Truth is never hidden: if many nodes are urgent, all may retain bright static borders.",
        "Motion is budgeted: at most one critical pulse and roughly two to three breathing active nodes in the same viewport.",
        "Selected project remains at full brightness; related projects may sit near 60%; unrelated projects may dim to 20–30% without disappearing.",
        "Selecting a node keeps its parents, children, dependencies, and resources legible while the distant graph dims.",
        "Temporal position answers when; glow answers how much deadline pressure. They are complementary, not duplicate channels.",
    ])

    s.h2("2.6 Zoom, project containers, and progressive disclosure")
    s.bullets([
        "Zoom out shows project containers; closer zoom reveals Deliverables; closer again reveals Work Packages and Actions.",
        "Project may be rendered as a group/container, not only a card. Zoomed-out Canvas shows the user's project landscape; zooming into a container reveals internal lineage.",
        "Zoom changes visual scale and semantic detail. Expand/collapse separately controls whether a specific node's children are exposed.",
        "Expand state persists per user and per view. A collapsed node reports hidden child count (for example, ‘6 hidden’).",
        "Proposal nodes must not look committed. Use translucency, partial glow, and dashed edges until Commit Split creates graph facts.",
    ])

    s.h2("2.7 Interaction details")
    s.bullets([
        "Hovering near the right edge/right corner of a block for roughly 0.8–1.0 seconds reveals a small plus affordance. It is hidden at rest to protect attention.",
        "The plus means Split / Extend and enters graph construction; it must not silently add an arbitrary todo.",
        "On mobile, replace hover with long-press or selection followed by an action menu.",
        "Node click opens the inspector or contextual detail; it does not turn the Canvas into an inline text editor.",
        "Primary edge routing is orthogonal: horizontal, vertical, and right-angle turns only. Avoid diagonal and arbitrary curved edges.",
    ])

    s.h2("2.8 Canvas controls and inspector contract")
    s.table(
        ["Surface", "Required controls / content"],
        [
            ["Global", "Project/Scope, Search, Filter, Zoom +/−, Fit, Auto Layout, Undo, and Canvas/Timeline/Mobile Preview switching."],
            ["Selected node", "Toggle, Split, Add Link/Dependency, Set Deadline, Resources, Status, Mark Done, Cancel, Supersede, Set as Current Project Scope, Move/reparent, and View History."],
            ["Inspector facts", "Title, stage/status, project/WBS level, description, deadline, estimated effort/span/finish, temporal pressure, progress, dependencies, children, resources, warnings, and history."],
            ["Status reasons", "DONE needs no reason. CANCELLED and SUPERSEDED show their required reason; Superseded may link to and open the replacement node."],
        ],
        [1800, 7560],
    )
    s.p("The first version should not overload double-click. Single click selects and updates the Inspector; expand/collapse and scope focus remain explicit controls. The current Project Scope may be an L1, L2, or L3 node according to the scale the user is managing.")

    s.h1("3. Timeline View — temporal graph")
    s.image(
        ASSETS / "timeline-source-reference.png",
        "Figure 2. Source reference supplied in the conversation: a contribution-style square activity grid.",
        "Small reference image showing daily square activity cells across months, used as inspiration for the Timeline block grid.",
        width=5.8,
    )
    s.image(
        ASSETS / "elastic-timeline-ui-concept-v2.png",
        "Figure 3. Original Elastic Timeline UI concept with module pool, day cells, overlap, capacity strip, and forecast panel.",
        "Dark Elastic Timeline desktop screen showing day-block scheduling, module cards, split-color overlap, deadline anchor, capacity strip, and non-judgmental forecast statistics.",
    )
    s.h2("3.1 Product philosophy")
    s.callout("Timeline is not a plan", "It is a Forecast + Capacity Map. Dragging work to a date means ‘I roughly intend this work to live around here,’ not ‘I must perform this task on that date.’")
    s.p("When work is not completed on the intended day, the forecast slides and recomputes. The system must not frame normal ADHD variability as failure.")

    s.h2("3.2 Two Timeline modes")
    s.table(
        ["Mode", "Question", "Form"],
        [
            ["Estimate", "Roughly how long will this project/module take, and will it meet the deadline?", "Compact temporal graph / Gantt-like summary at Project or Module level; dependency and deadline remain visible."],
            ["Block", "Where does this work occupy calendar capacity, overlap with other work, or overload a period?", "Left card pool plus right square day-block grid; draggable/resizable regions; overlap split by project color."],
        ],
        [1500, 3300, 4560],
    )

    s.h2("3.3 Block mode layout and behavior")
    s.bullets([
        "Left panel holds unscheduled Project/Module/Deliverable/Work Package cards with estimated effort, deadline, and tags.",
        "Right panel is a horizontal time grid. The first version uses one square per day; day/week/month zoom may follow.",
        "Timeline primarily schedules WBS L2–L3. Atomic Actions support estimation but are not all rendered as draggable blocks by default.",
        "A scheduled region spans several day cells and displays the project/module name across the region.",
        "When two projects occupy a day, the cell uses a clean half-and-half split (left/right is the preferred first implementation; top/bottom is an alternative).",
        "A module may be resized by its right edge. The forecasted length may initialize the region, but user movement remains a rough intention.",
        "Explicit deadline is a diamond/anchor and creates a forbidden region. Children cannot be dragged beyond the relevant parent deadline.",
        "A Capacity Strip below the grid shows daily load density and identifies overload without blaming the user.",
    ])

    s.h2("3.4 Elastic forecast model")
    s.p("The accepted model estimates delivery from completed workload, not timers and not an assumed fixed number of work hours per day.")
    s.code("estimated_effort = task size in hour-equivalent\ncompleted_estimated_effort = sum(estimated_effort for user-marked DONE actions)\nobserved_delivery_pace = completed_estimated_effort / calendar window\ntypical_pace = median weekly pace over the most recent 4–8 weeks\nestimated_calendar_span ≈ project_estimated_workload / typical_pace\nestimated_finish = forecast range adjusted for dependencies, existing project load, and deadline constraints")
    s.bullets([
        "Observed Delivery Pace is not actual labor time. It means the amount of estimated workload the user has historically closed in real calendar time.",
        "Use the most recent 4–8 weeks. Typical pace is approximately the median; variability should remain visible rather than being averaged away. The exact variability statistic was not fixed in the conversation.",
        "Display a likely range, not false precision. Example: typical span about 5 weeks; likely range 4–7 weeks.",
        "First version may use one overall pace. Task-type-specific pace (Read, Write, Coding, and so on) is a future enhancement.",
        "Graph changes trigger derived-state recomputation: workload, estimated span, estimated finish, deadline gap, temporal pressure, and Canvas glow.",
        "Estimated Finish and Deadline are separate fields. Their gap creates a warning, for example Estimated Finish 12 Oct versus Deadline 30 Sep.",
    ])
    s.code("Example\nDONE over 14 calendar days: 1h + 2h + 3h + 1h = 7h-equivalent\nObserved Delivery Pace = 3.5h-equivalent/week\nNew project workload = 20h; typical pace ≈ 4h/week\nForecast: about 5 weeks, likely range 4–7 weeks")

    s.h2("3.5 Effort, observed time, and task-type defaults")
    s.bullets([
        "Estimated effort describes task size. It is not a promise about duration or a daily plan.",
        "Estimate source precedence: user value when provided; otherwise AI suggestion; otherwise Task Type default.",
        "Source examples for configurable defaults: Read 30m, Search 45m, Research 60m, Write 90m, Design 120m, Build 120m, Discussion 30m. These are starting values, not universal truths.",
        "Do not ask a completion Q&A such as ‘How long did it take?’ and do not require a timer.",
        "Optional passive observation may record observed_work_time from active sessions, but this is only system-observed time and must never be labeled actual effort.",
        "The first version does not depend on observed_work_time. Delivery pace comes from user-marked DONE workload.",
    ])

    s.h2("3.6 Timeline and Canvas relationship")
    s.table(
        ["Dimension", "Canvas projection", "Timeline projection"],
        [
            ["Primary emphasis", "Hierarchy, relation, state, resource", "Time, overlap, load, forecast, deadline"],
            ["Temporal pressure", "Whole-card glow", "Explained score/range and capacity context"],
            ["Deadline", "Anchor/constraint on the spatial graph", "Visible date anchor and forbidden region"],
            ["Expand state", "Canvas-specific", "Timeline-specific"],
            ["Atomic Actions", "Visible on demand", "Usually hidden; used to calculate Module length"],
        ],
        [1800, 3780, 3780],
    )

    s.h2("3.7 Timeline controls and block actions")
    s.table(
        ["Surface", "Required controls / behavior"],
        [
            ["Timeline global", "Estimate, Block Grid, Today, Previous, Next, Day/Week/Month/Quarter, Filter, Unscheduled, Fit Project, Forecast, and Toggle."],
            ["Scheduled block", "Open, Toggle, Set Deadline, Resize Span, Remove from Timeline, View in Canvas, and Resources."],
            ["Remove from Timeline", "Removes only the rough temporal placement. It must never delete the graph node."],
            ["Cross-view locate", "A selected node should be easy to locate in either projection through View in Canvas / View in Timeline."],
        ],
        [1900, 7460],
    )

    s.h1("4. Cross-view data and view state")
    s.bullets([
        "Graph facts include node identity, hierarchy, dependencies, lifecycle, deadlines, estimates, tags, and resource references.",
        "View state includes per-view expanded nodes, selected scope, filters, zoom/pan, and optional manual placement. It must not pollute graph facts.",
        "Proposal state is separate from committed graph state.",
        "Canvas and Timeline must be projections of the same graph; they may not maintain divergent task records.",
        "Notion and future views may read the same source of truth without changing the core schema.",
    ])

    s.h1("5. ADHD and INTP positioning")
    s.p("The product should not claim to solve or treat ADHD. Its design intent is to reduce executive-function load by externalizing goal hierarchy, dependency, deadline, next-step clarity, status, and output location. Likewise, INTP/MBTI language may describe preferences but should not be presented as scientific proof of interface suitability.")
    s.bullets([
        "Canvas serves users who want system-level understanding and logic.",
        "Focus/mobile serves users who need the universe hidden while acting.",
        "Progressive disclosure, quiet motion, and persistent toggles reduce visual overload.",
        "The elastic Timeline accommodates bursty, inconsistent real-life delivery rather than enforcing calendar discipline.",
        "The Inbox/Capture concept remains important: users may capture an idea without immediately deciding its project or WBS level; human and AI can place it later.",
    ])

    s.h1("6. Decision ledger")
    s.table(
        ["Status", "Decision"],
        [
            ["Accepted", "Canvas is a constrained graph dashboard with orthogonal edges, project hue, WBS lightness, perimeter lifecycle, Timeliner-derived glow, warning icon, persistent toggle, auto-layout-first behavior, and desktop emphasis."],
            ["Accepted", "Dedicated Timeline View is separate from Canvas and contains Estimate and Block modes; Block mode uses a left card pool, daily square cells, split-color overlap, deadline anchors, capacity strip, and elastic forecast."],
            ["Accepted", "Timeline is forecast/capacity, not a fixed plan. Observed Delivery Pace over recent 4–8 weeks drives the first version; no timer or completion Q&A."],
            ["Accepted", "Mobile avoids complex dragging; use Focus/Current Project/Action cards and long-press or selection menus."],
            ["Corrected", "Status is expressed by whole-card perimeter and aura, not only a right-corner light."],
            ["Deferred", "Exact variance statistic, task-type-specific pace, day/week/month detail, and optional character/emblem styling."],
            ["Rejected / superseded", "A single generic attention_level independent of time; glow is primarily Temporal Pressure. A rigid daily capacity plan is not the forecast model."],
            ["Out of scope for this project system", "Agent Work Queue and Human Review Queue management; another orchestration system may own them."],
        ],
        [1800, 7560],
    )

    s.h1("7. Reference reconciliation and MVP acceptance checklist")
    s.p("Reference coverage: the supplied v0.1 specification’s product principle and visual-encoding rules are represented in Sections 1–2; Canvas sections 20–39 in Section 2; Timeline sections 40–52 in Section 3; Mobile, Portfolio, and Inbox sections 53–55 in Sections 1 and 5; button and non-violation summaries in Sections 2.8, 3.7, 6, and this checklist. This mapping is included so compact synthesis does not silently erase source requirements.")
    s.bullets([
        "Canvas renders project containers and WBS nodes from the graph source; edges are orthogonal and auto-layout is available.",
        "Project hue and WBS lightness are deterministic; lifecycle borders and temporal-pressure glow never conflict with project color.",
        "Anchored deadlines and child/dependency drag constraints are enforced with visible forbidden regions.",
        "Expand/collapse persists separately for Canvas and Timeline; proposal nodes are visually distinct from committed nodes.",
        "The edge-hover plus appears only after deliberate hover and enters a Split Session.",
        "Timeline Block mode schedules L2/L3 cards in daily squares, supports resize, renders overlap, and shows deadlines/capacity.",
        "Timeline forecast uses estimates plus 4–8-week Observed Delivery Pace and returns a range, Estimated Finish, Deadline, and gap warning.",
        "No completion-time Q&A or mandatory timer is introduced.",
        "The mobile view remains readable and action-oriented without graph dragging.",
    ])

    s.sources(extra=[("FlowTune dashboard visual reference", SOURCE_UI_REFERENCE)])
    return s.save("01_Canvas_and_Elastic_Timeline_UI_Spec.docx")


def build_graph_spec() -> Path:
    s = SpecDoc(
        "Graph-Based Project Management System",
        "Graph core, runtime boundaries, data model, operations, completion semantics, and delivery forecast",
        "Graph-Based Project System",
        accent=CYAN,
    )
    s.add_cover("System specification v1.0 — 28 Aug 2026", "Architecture, product, data, agent integration")

    s.h1("Information preservation note")
    s.p("This specification preserves the full conceptual progression in the source discussion and reconciles it with Yonc_Graph_Project_System_Spec_v0.1.md: general Graph Engineering, the existing Yonc repository baseline, the proposed graph architecture, later scope reductions, and the accepted completion/timeline decisions. Earlier ideas that were superseded remain documented as history rather than being silently merged into the final model.")
    s.callout("System definition", "A human–agent collaborative project graph that stores project truth, structure, time, state, and resource references; supports reasoning and graph operations; and projects the same source into Canvas, Timeline, Portfolio, Mobile, and optional Notion views.")

    s.h1("1. Graph Engineering context")
    s.p("Graph Engineering is the explicit design of nodes, edges, shared state, routing, branches, loops, gates, parallelism, verification, permissions, recovery, and stop conditions. It is not merely drawing a graph and it does not replace agent loops.")
    s.table(
        ["Concept", "Meaning in an agent/project system"],
        [
            ["Node", "Agent, tool, function, validator, human approval, project item, or promoted artifact."],
            ["Edge", "The explicit relationship or next path: contains, depends_on, blocks, produces, validates, and others."],
            ["State", "Shared project and runtime facts carried between operations."],
            ["Branch", "Conditional path selected from current state or validation result."],
            ["Loop", "Repeated work/check/fix cycle inside a node or agent."],
            ["Gate", "Condition or approval required before transition."],
            ["Parallel", "Independent nodes or agents that may proceed concurrently."],
        ],
        [1700, 7660],
    )
    s.p("Graph definitions may be expressed in JSON, YAML, code-based builder APIs, or database rows. The format is secondary; the essential structure is nodes + edges + conditions + runtime/state behavior. For a Hermes-style multi-agent platform, graph logic belongs in the orchestration layer while each agent can keep its own internal plan/act/check loop.")
    s.code("Prompt → Context → Harness → Loop → Graph → Agent System\n\nGraph = coordination among loops/agents/tools\nLoop  = local plan → act → check → fix")

    s.h1("2. Existing Yonc baseline from the source discussion")
    s.p("The source discussion reported that the repository already contains dashboard/state/flow foundations rather than a blank agent shell. These notes are preserved as the design baseline and should be revalidated against the current code during implementation.")
    s.bullets([
        "dashboard.py, current_state.json, flow_runs logs, cron/state/config readers, and Timeliner-related logic already exist.",
        "README describes a system for INTP + ADHD that decomposes abstract goals to executable action.",
        "Existing WBS: L1 Goal → L2 Deliverable → L3 Work Package → L4 Activity → Atomic Refinement; Project and Exploratory OKR use different decomposition patterns.",
        "task_reader.py reportedly builds recursive records with id, title, parent_id, depth, and children. This is already a hierarchy tree even if Notion displays bullets.",
        "dashboard.py reportedly groups tasks by Mode, Task Type, and Theme, answering ‘what should I see today?’ but not ‘what does the entire project look like?’",
        "state_evaluator.py reportedly derives lifecycle from node properties, and main flow already calls the evaluator while older flow-l1/l2/l3 paths are legacy wrappers.",
        "Timeliner rank/scope already affects ordering and flow. Timeliner should therefore remain a time/pressure engine rather than being replaced by the new graph manager.",
    ])
    s.table(
        ["Existing-state sequence reported in source", "Interpretation"],
        [
            ["RAW → STRUCTURED → SCOPED → SEQUENCED", "A node acquires structure, scope, and ordering."],
            ["EXPANDING → HUMAN_REVIEW", "The node is split and the AI proposal awaits human approval."],
            ["PHASING_WAIT → ACTIONABLE_PENDING", "Timing/phase or Mode/Task Type details remain unresolved."],
            ["READY → COMPLETED", "The node becomes executable and later complete."],
        ],
        [4200, 5160],
    )
    s.p("The later product discussion introduced a different lifecycle vocabulary (raw, clarified, structured, split needed, split proposed, approved, ready, doing, review, done) plus stage/status separation. A final one-to-one migration mapping was not resolved and remains an implementation task.")

    s.h1("3. System boundary and source of truth")
    s.bullets([
        "Notion is one UI/projection. The source of truth is the repository-backed graph store, initially compatible with an upgraded current_state.json or equivalent graph JSON.",
        "All views read the same graph. Canvas, Timeline, Mobile, Portfolio, Notion, and external agents must not maintain divergent task truth.",
        "The system owns project truth, structure, lifecycle, deadlines, estimates, dependencies, resources/artifacts, warnings, and history.",
        "The project system does not own general Agent Work Queue or Human Review Queue orchestration. That was explicitly removed from scope.",
        "External agents may read the graph, propose work, attach outputs, or request permitted updates; user completion authority remains final.",
    ])
    s.code("Project Graph scope\n= Project truth + structure + timeline + visualization + operations\n≠ universal agent scheduler / workload queue")

    s.h1("4. Logical architecture")
    s.table(
        ["Component", "Responsibility"],
        [
            ["Graph Store", "Authoritative graph ID/version, committed nodes/edges, resource references, derived values, and view-independent facts."],
            ["Graph Builder", "Transforms a fuzzy goal or approved Split Session proposal into candidate nodes and edges."],
            ["Graph Evaluator", "Checks graph health: actionability, missing outputs, orphan nodes, dependency cycles, deadline conflicts, and other constraints."],
            ["Graph Operations", "Only supported mutation path: create/link/split/commit/move/set deadline/mark done/cancel/supersede/attach resource; produces event history and undo."],
            ["Graph Runtime", "Evaluates state transitions, route/edge conditions, readiness, derived values, and recomputation. General agent scheduling is outside this product's MVP scope."],
            ["Timeliner", "Owns temporal rank, deadlines, estimated span/finish, delivery pace, pressure, and time-related constraints."],
            ["View Projector", "Projects the same graph into Canvas, Timeline, Portfolio, Mobile/Focus, and optional Notion views."],
            ["Event History", "Records operations needed for review, recovery, sync, and undo. Detailed changed_by metadata is useful but was downgraded from product core."],
        ],
        [2100, 7260],
    )
    s.code("User / Agent\n      ↓\nGraph Builder / Split Session\n      ↓ candidate\nDraft Proposal Graph\n      ↓ user Commit\nGraph Operations → Graph Store → Evaluator / Runtime / Timeliner\n                                  ↓\n                      Canvas | Timeline | Mobile | Notion")

    s.h1("5. Graph data model")
    s.h2("5.1 Top-level graph")
    s.code("graph\n├─ id\n├─ version\n├─ nodes[]\n├─ edges[]\n├─ derived_state / runtime data\n└─ view_state references (stored separately from graph facts)")

    s.h2("5.2 Node core")
    s.table(
        ["Field group", "Fields / semantics"],
        [
            ["Identity", "id, title, node_type, wbs_level, project/theme"],
            ["Classification", "tags, Mode, Task Type, priority, phase where applicable"],
            ["Lifecycle", "stage and status as separate concepts; completion remains an explicit status fact"],
            ["Time", "planned/rough start, deadline, estimated_effort, estimated_span, estimated_finish, optional observed_work_time, temporal_pressure"],
            ["Execution definition", "start cue, inputs, done_when/output definition; often stored inside an Action rather than as more graph nodes"],
            ["Resources", "resource references; important resources may be promoted to Artifact nodes"],
            ["Metadata", "confidence, optional change/audit details, generated/human review markers, other extensible properties"],
        ],
        [2100, 7260],
    )
    s.p("Core node types begin with project/goal, deliverable/module, work_package, and action. Artifact, Resource, and Agent may become graph nodes only when their relationships justify it. Simple files and links remain references.")

    s.h2("5.3 Edge taxonomy")
    s.table(
        ["Category", "Edge types", "Notes"],
        [
            ["Structural", "contains", "Carries required=true/false for progress and closure calculations."],
            ["Logical", "depends_on, blocks, blocked_by, related_to, validates", "Makes ordering, blocking, verification, and cross-linking explicit."],
            ["Execution", "executed_by, uses", "Connects work to agents/tools/resources when needed; assignment does not imply the graph system schedules a global queue."],
            ["Output", "produces, used_by", "Links work to artifact/resource references or promoted artifact nodes."],
            ["History", "superseded_by", "Preserves replacement history without deleting the prior node."],
        ],
        [1700, 2800, 4860],
    )
    s.p("The accepted minimal edge set for an initial graph upgrade was contains, dependency/depends_on, and blocks. The broader taxonomy is retained for system evolution.")

    s.h2("5.4 Resource and artifact model")
    s.bullets([
        "Artifact is first an Artifact Reference / pathway, not necessarily an uploaded file or a graph node.",
        "Supported reference styles may include local://, github://, notion://, https://, s3://, drive://, and conversation://.",
        "Promote a resource to an Artifact Node when it has versions, is shared by many tasks, produces other artifacts, or needs explicit graph relationships.",
        "The graph needs to know where an output lives; it does not need to become a full file manager.",
    ])

    s.h1("6. Draft, committed graph, operations, and view state")
    s.bullets([
        "Draft/Proposal Graph is separate from Committed Graph. AI-proposed nodes and edges do not become project truth until user Commit.",
        "Every change is a graph operation rather than an arbitrary JSON edit. Examples: propose_split, approve/commit_split, link_dependency, set_deadline, move_schedule, mark_done, cancel, supersede, attach_resource.",
        "Committing a split emits CREATE_NODE and CREATE_EDGE operations; the event history enables review and undo.",
        "View state (expanded/collapsed, selected scope, filters, zoom/pan, manual layout) is separate from graph facts and may vary by view.",
        "UI placement must not be hard-coded into task truth. A future UI can project the same graph differently.",
    ])

    s.h1("7. Lifecycle, status, and completion authority")
    s.h2("7.1 Stage versus status")
    s.p("Stage describes the lifecycle phase; status/runtime state describes what is happening now. A node may be stage=execution and status=blocked, or stage=planning and status=proposed. The exact normalized enums must be reconciled with the existing state_evaluator vocabulary.")

    s.h2("7.2 Completion rules")
    s.bullets([
        "DONE has one final rule: the user explicitly says/marks Done. This applies to Action, Module, and Parent nodes.",
        "Agents may prepare work, produce artifacts, submit, or propose completion, but they do not unilaterally set DONE by default.",
        "Children 4/4 complete does not auto-complete the Parent. It is a calculated fact only; the user may add validation or declare the scope complete.",
        "The previously proposed READY_TO_CLOSE / closure-suggestion flow was superseded by the simpler explicit-user-Done rule. A non-blocking suggestion may exist, but it is not a status transition requirement.",
        "DONE does not require a reason because that would add friction.",
        "CANCELLED requires reason. SUPERSEDED requires reason and may link superseded_by to the replacement node.",
    ])

    s.h2("7.3 Progress calculation")
    s.code("progress = completed estimated effort of required Actions\n           / total estimated effort of required Actions\n\nFallback when estimates are absent:\ncompleted required children / total required children")
    s.bullets([
        "Progress and status are independent. A node may be user-marked DONE while calculated progress remains 87% because some planned actions became unnecessary.",
        "Only required=true children contribute to default progress. Optional work does not prevent 100%.",
        "Task-count progress is a fallback because a 15-hour action must not equal a 20-minute action.",
    ])

    s.h1("8. Time, effort, forecast, and recomputation")
    s.table(
        ["Field", "Meaning"],
        [
            ["estimated_effort", "Relative task size in hour-equivalent; user value > AI suggestion > Task Type default."],
            ["observed_work_time", "Optional passively observed active-session time; never claimed as actual labor time."],
            ["observed_delivery_pace", "Completed estimated workload per calendar window; accepted first-version historical signal."],
            ["estimated_span", "Likely natural-calendar range; not a fixed plan."],
            ["estimated_finish", "Dynamic forecast range derived from workload, delivery pace, dependency, current load, and deadline constraints."],
            ["deadline", "Hard anchor the project must not cross; separate from estimated finish."],
            ["temporal_pressure", "Timeliner-derived pressure shown as Canvas glow and explained in Timeline."],
        ],
        [2500, 6860],
    )
    s.p("Graph mutations must propagate. Adding an action, changing effort, completing work, changing a dependency, or moving a deadline recomputes parent workload, estimated span/finish, deadline gap, temporal pressure, and Canvas/Timeline projections.")
    s.bullets([
        "Use the recent 4–8 weeks of user-marked DONE estimated effort to calculate typical delivery pace.",
        "Use a typical/median pace and preserve variability; show ranges rather than false precision.",
        "Task-type-specific pace and estimate-confidence visualization remain future enhancements.",
        "No mandatory timer and no completion-time Q&A.",
    ])

    s.h1("9. Graph health and constraints")
    s.bullets([
        "Child beyond parent deadline",
        "B depends on A but is planned before A",
        "Circular dependency",
        "Orphan node",
        "Large node that remains unsplit / fails the Actionability Contract",
        "Action missing an output or done_when definition",
        "Blocked too long",
        "Agent output stalled or submitted without a resource reference (future integration)",
    ])
    s.p("UI behavior is intentionally quiet: normally show one small warning icon in the card corner; clicking reveals reason and affected nodes. Health checks advise and expose structural issues rather than making hidden decisions for the user.")

    s.h1("10. Agent and tool interface")
    s.p("The Project Graph acts as a shared coordination substrate. Main Orchestrator and external agents may query it, open a Split Session, propose execution, link outputs, and submit permitted operations. A Coding, Research, BIM, or Reviewer agent can keep its own internal loop while the outer graph defines relationships and gates.")
    s.code("project.list / project.get_graph / project.get_node\nproject.get_children / project.get_dependencies / project.get_resources\nsplit.start / split.get_context / split.propose / split.revise\nsplit.get_current_proposal / split.commit\nnode.mark_done / node.cancel / node.supersede\nnode.set_deadline / node.add_resource")
    s.p("These names are an interface direction, not a frozen API. Final permissions must ensure agents cannot silently rewrite committed project truth or completion status.")

    s.h2("10.1 End-to-end processing flow")
    s.numbered([
        "User creates or imports a project and the initial graph node exists.",
        "User opens a Split Session; user and AI discuss and agree on decomposition.",
        "User commits the split; graph operations create children and edges.",
        "Estimated effort is assigned from user input, AI suggestion, or Task Type default.",
        "Derived state recalculates; Canvas re-renders and Timeline forecast updates.",
        "User may place modules roughly on the Timeline; this changes temporal intent, not graph identity.",
        "The user works in real life or external agents produce artifacts; resource references attach to the relevant nodes.",
        "User explicitly marks DONE; Observed Delivery Pace updates and future forecasts become progressively more realistic.",
    ])

    s.h1("11. MVP and phased evolution")
    s.table(
        ["Phase", "Scope"],
        [
            ["MVP", "Graph source of truth; core nodes; contains/depends_on/blocks; operations/history; collaborative split; user-only Done; resources; Canvas and Timeline projections; delivery-pace forecast."],
            ["Next", "Graph health expansion, promoted artifact nodes, richer external-agent read/propose interfaces, uncertainty ranges, task-type delivery pace."],
            ["Later", "Conditional multi-agent graph runtime, richer checkpointer/scheduler integration, dynamic graph generation, advanced validation and permissions."],
        ],
        [1500, 7860],
    )

    s.h1("12. Decision ledger")
    s.table(
        ["Status", "Decision"],
        [
            ["Accepted", "Graph Engineering means nodes + edges + rules + state transitions + runtime reasoning; it is not merely graph storage."],
            ["Accepted", "JSON/repository graph is source of truth; Notion and all other UIs are projections."],
            ["Accepted", "Graph Store, Builder, Evaluator, Operations, Runtime, Timeliner, View Projector, and Event History form the logical architecture."],
            ["Accepted", "Draft proposals stay separate until Commit; view state stays separate from graph facts; operations/history support undo."],
            ["Accepted", "User alone defines DONE; required/optional supports weighted progress; Cancelled/Superseded require reason."],
            ["Accepted", "Resource references are lightweight pathways; promote to Artifact Node only when relationships require it."],
            ["Deferred", "Exact lifecycle migration mapping, estimate confidence UI, task-type-specific pace, detailed change-source metadata."],
            ["Out of scope", "Owning general agent work queues and human review queues; arbitrary autonomous scheduling across agents."],
        ],
        [1800, 7560],
    )

    s.p("Reference coverage: the v0.1 specification’s product definition/scope, source of truth, node/edge models, actionability, operations, completion, estimation, derived state, resources, external-agent interface, architecture, processing flow, non-violation rules, and MVP are preserved in Sections 1–12. Canvas/Timeline interaction detail is intentionally routed to Document 01; the full collaborative split workflow is routed to Document 03.")

    s.sources()
    return s.save("02_Graph_Based_Project_Management_System_Spec.docx")


def build_split_spec() -> Path:
    s = SpecDoc(
        "Collaborative Task Decomposition",
        "Split Session workflow, Actionability Contract, proposal lifecycle, and commit semantics",
        "Collaborative Task Decomposition",
        accent=GOLD,
    )
    s.add_cover("Workflow specification v1.0 — 28 Aug 2026", "Product, agent design, UX, graph operations")

    s.h1("Information preservation note")
    s.p("This document isolates every substantive task-decomposition decision from the source conversation and Yonc_Graph_Project_System_Spec_v0.1.md: the LineV2/WBS hierarchy, OKR awareness, the rejection of black-box auto-splitting, the full Split Session workflow, stop conditions, bounded exploration, external-agent/MCP entry points, proposal/commit separation, completion authority, and the accepted estimate defaults.")
    s.callout("Non-negotiable rule", "Task decomposition is a negotiated human–AI workflow. The agent proposes and revises; the user decides whether the graph is actionable enough and explicitly commits the split.")

    s.h1("1. Purpose and non-goals")
    s.bullets([
        "Turn a fuzzy Goal, Objective, Key Result, Deliverable, Module, or Work Package into a small set of actionable graph nodes.",
        "Reduce decision load at the point of execution without exploding the graph into meaningless micro-steps.",
        "Preserve the parent intent, output expectations, dependencies, required/optional status, estimates, tags, and resource context.",
        "Keep AI-proposed decomposition outside committed project truth until user Commit.",
        "Do not act as a black-box automatic splitter and do not pretend uncertain work has a fully known step sequence.",
    ])

    s.h1("2. Structural model: Lines, OKR, WBS, and project scope")
    s.table(
        ["Level / concept", "Role"],
        [
            ["Line / Programme / Objective", "Long-running direction or major outcome; may contain Key Results or major projects."],
            ["L1 Goal", "High-level desired state; usually too large to execute directly."],
            ["L2 Deliverable / Module", "Substantial project-shaped result; may appear as a Portfolio project card."],
            ["L3 Work Package", "Coherent cluster of work that produces part of a Deliverable."],
            ["L4 Action", "Directly startable work that satisfies the Actionability Contract."],
            ["Atomic Refinement", "Lightweight execution definition inside an Action; not automatically another graph layer."],
        ],
        [2400, 6960],
    )
    s.p("The existing Yonc design reportedly supports both Project and Exploratory OKR decomposition. The new Split capability must therefore be OKR-aware: Objective → Key Results → graph work, while retaining the WBS lineage used by LineV2. WBS level alone never proves that a node is actionable.")
    s.bullets([
        "Any node may be chosen as current Project Scope, making decomposition relative to the user's current scale.",
        "A WBS L2 Module can be managed as a project without changing its underlying graph identity.",
        "Inbox/Capture remains a valid pre-decomposition layer: an idea may be captured before the user knows its project or WBS level.",
    ])

    s.h1("3. Actionability Contract")
    s.callout("Action definition", "An Action can be started without another round of ‘how do I do this?’ planning and has an observable result that lets the user decide whether it is complete.")
    s.table(
        ["Criterion", "Test"],
        [
            ["Startable", "If shown now, does the user know the first physical/digital step?"],
            ["Observable Done", "Is there a concrete output, state change, or result that makes completion discussable?"],
            ["Single Intent", "Is it one main action rather than research + compare + write + revise hidden inside one label?"],
            ["Decision Load", "Does execution still require a major planning/design decision? If so, split or clarify further."],
            ["Bounded Effort", "Can it advance in one or a small number of reasonable work sessions? This is a soft rule, not a strict two-hour threshold."],
        ],
        [1800, 7560],
    )
    s.p("Decision Load is the most important ADHD-oriented test. A short-looking task may still be bad if it hides many choices. Conversely, a 2.5-hour task can be a good Action if its input, intended output, and method are already clear.")
    s.code("Bad:  ‘Write Methodology’\nBetter: ‘Using prototype-test.md, turn the existing experiment steps into a 400–600 word Methodology 2.3 draft.’")

    s.h2("3.1 Atomic Refinement inside an Action")
    s.p("Stop expanding the graph at the Action by default. Store the final activation details inside the node.")
    s.code("Action: Write Graph Core design spec\nStart: Open graph-core.md\nInputs: Previous architecture discussion\nDone when:\n  - Node model documented\n  - Edge model documented\n  - One architecture diagram exists")
    s.p("This preserves a manageable Canvas while still lowering activation energy.")

    s.h2("3.2 Bounded exploration")
    s.p("Unknown work such as debugging or research cannot always be honestly decomposed into deterministic steps. Convert it into a bounded exploration whose next action and output are known even though the answer is not.")
    s.code("Investigate intermittent crash for 45 minutes\nOutput:\n  - Reproduction steps\n  - Three likely causes\n  - Next experiment")
    s.p("The key distinction is: the answer may be uncertain; the next action must not be.")

    s.h1("4. Split Session workflow")
    s.numbered([
        "Open a Split Session from a selected parent node.",
        "Load parent goal, lineage, existing children, dependencies, resources, tags, lifecycle, deadline, estimates, and relevant constraints.",
        "Agent identifies ambiguity and proposes decomposition version 1 with rationale, outputs, dependencies, estimates, and required/optional flags.",
        "User responds conversationally: merge, remove, rename, keep coarse, change output, change sequence, or split further.",
        "Agent revises the proposal and keeps a visible current proposal alongside the chat.",
        "Agent evaluates the proposed leaves against the Actionability Contract and explains why it believes each is actionable.",
        "User chooses Continue splitting, Modify, or Commit. Consensus is practical agreement, not an autonomous AI verdict.",
        "Commit emits graph operations that create nodes/edges and records history for review and undo.",
    ])
    s.code("OPEN SESSION → Proposal v1 → User feedback → Proposal v2 → …\n→ Actionability check → User Commit → CREATE_NODE / CREATE_EDGE operations")

    s.h1("5. Entry points and reusable service")
    s.p("The Canvas sub-chat is one entry point, not the implementation boundary. The same workflow must be callable from desktop, mobile, an external chat, another agent, or an MCP/tool interface.")
    s.table(
        ["Entry point", "Behavior"],
        [
            ["Canvas", "Hover right edge/right corner about 0.8–1.0 seconds to reveal +; click opens compact split sub-chat and proposal panel."],
            ["Mobile", "Long-press or select a node, then choose Split; use a full-screen conversational flow rather than graph dragging."],
            ["External agent/chat", "Agent identifies graph/node, starts a session, discusses with the user, then calls commit only after explicit approval."],
            ["MCP/tool API", "Programmatic access to session context, proposals, revision, validation, and commit."],
        ],
        [1900, 7460],
    )
    s.code("split.start(parent_node_id)\nsplit.get_context(session_id)\nsplit.propose(session_id, proposal)\nsplit.discuss(session_id, user_feedback)\nsplit.revise(session_id)\nsplit.get_proposal(session_id)\nsplit.commit(session_id)")

    s.h1("6. Split Session UI")
    s.table(
        ["Area", "Required behavior"],
        [
            ["Session header", "Parent node title, lineage, lifecycle, deadline, and why splitting was invoked."],
            ["Conversation", "AI rationale and user feedback in a compact chat; retain context without turning the Canvas into a chat wall."],
            ["Current proposal", "Live tree/graph preview with outputs, required/optional flags, dependencies, estimates, and validation indicators."],
            ["Proposal status", "Clearly draft; semi-transparent/partial border/dashed edges; never visually equal to committed graph."],
            ["Actions", "Continue discussion, Split deeper, Modify, Discard, Commit Split."],
        ],
        [1900, 7460],
    )
    s.p("The plus affordance means enter Graph Construction. It is not a quick-add todo button. The user must always be able to see the current proposal before commit.")

    s.h1("7. Proposal and session data")
    s.table(
        ["Object", "Minimum fields"],
        [
            ["SplitSession", "id, parent_node_id, state, context_snapshot/version, created_at, current_proposal_id"],
            ["Proposal", "id, version, rationale, proposed_nodes, proposed_edges, actionability_results, warnings"],
            ["Proposed node", "title, node_type/WBS, output/done_when, estimate, tags, required/optional, uncertainty, resource inputs"],
            ["Proposed edge", "type, source, target, required flag, rationale/condition where needed"],
            ["Commit result", "operation IDs, created node/edge IDs, graph version, warnings, undo reference"],
        ],
        [2100, 7260],
    )
    s.p("Suggested session states: OPEN, PROPOSAL_DRAFT, PENDING_USER_REVIEW, REVISING, APPROVED, COMMITTED, DISCARDED. The exact enum names were not frozen in the source; the essential rule is that proposed and committed states remain distinct.")

    s.h1("8. Estimate behavior during decomposition")
    s.bullets([
        "Each Action receives estimated_effort. Source precedence is user value, AI suggestion, then Task Type default when blank.",
        "Accepted configurable default examples: Read 30m, Search 45m, Research 60m, Write 90m, Design 120m, Build 120m, Discussion 30m.",
        "AI may offer a range (for example 45–90 minutes) with an editable working estimate.",
        "The estimate describes task size; it is not a calendar promise.",
        "Do not ask the user to report actual effort on completion. No mandatory timer or completion Q&A.",
        "User-marked DONE effort later contributes to Observed Delivery Pace and Elastic Timeline forecasting.",
    ])

    s.h1("9. Commit, graph operations, history, and undo")
    s.bullets([
        "Commit validates that the parent graph version still matches the session context or explicitly resolves conflicts.",
        "Create all proposed child nodes and structural edges as a coherent operation batch; add dependency/output edges in the same batch where appropriate.",
        "Preserve required/optional semantics on contains edges.",
        "Record the committed proposal version and graph operation IDs so the change can be reviewed and undone.",
        "A split revision that is not committed does not alter project truth.",
        "Graph health re-runs after commit. New work triggers workload, estimate, deadline-gap, and temporal-pressure recomputation.",
    ])

    s.h1("10. Completion and agent participation")
    s.bullets([
        "Agents may inspect the graph, volunteer for agent-capable work, prepare artifacts, and submit results, subject to permission.",
        "The project system stores task structure/state and artifact references; general agent work/human review queues remain out of scope.",
        "Agent completion is a proposal/submission. DONE occurs only when the user explicitly says/marks Done.",
        "Children all complete does not auto-complete the Parent. Calculated progress and status are separate.",
        "A simple resource may be linked as a pathway; a highly connected result may be promoted to an Artifact Node.",
    ])

    s.h1("11. Worked examples")
    s.h2("11.1 Product module")
    s.code("Parent: Develop Project Management Module\n\nProposal:\n1. Define project data model\n2. Build Project API (depends on 1)\n3. Build Project Map UI (depends on 1)\n4. Integrate graph operations (depends on 2 + 3)\n5. Test project workflow (depends on 4)")
    s.p("The proposal remains draft while discussed. After Commit, steps 2 and 3 become parallel-ready when step 1 is user-marked Done; the graph can then unlock step 4.")

    s.h2("11.2 Research gap")
    s.code("Too large: Research Gap\nSuggested decomposition:\n- Search papers\n- Extract theory\n- Extract results\n- Extract limitations\n- Build gap matrix\n- Define gap\n\nDependency idea:\nSearch → parallel extraction → Gap matrix → Gap definition")

    s.h2("11.3 Actionability refinement")
    s.code("Not actionable: Research Graph Engineering\nActionable: Find five agent-graph architecture implementations and record each implementation's node, edge, and runtime approach in graph-notes.md.")

    s.h1("12. Graph-health checks related to splitting")
    s.bullets([
        "Node is too large or repeatedly fails Actionability criteria.",
        "Leaf lacks a clear output or done_when definition.",
        "Proposal creates orphan nodes, circular dependencies, or an impossible order.",
        "Child estimate/deadline makes the parent forecast exceed its deadline.",
        "Proposal duplicates existing children or artifacts.",
        "Proposal decomposes an uncertain problem into false certainty instead of bounded exploration.",
        "Graph expansion creates too many trivial nodes; recommend Atomic Refinement inside an Action.",
    ])

    s.h1("13. Decision ledger")
    s.table(
        ["Status", "Decision"],
        [
            ["Accepted", "Decomposition is a Collaborative Split Session with multiple proposal/revision rounds and explicit user Commit."],
            ["Accepted", "Split capability is service-level and reusable from Canvas, Mobile, external chat/agent, and MCP/tool interfaces."],
            ["Accepted", "Actionability Contract uses Startable, Observable Done, Single Intent, Decision Load, and Bounded Effort; duration alone is insufficient."],
            ["Accepted", "Graph stops at Action by default; Atomic Refinement provides start/inputs/done_when without producing graph clutter."],
            ["Accepted", "Bounded exploration handles genuinely uncertain work without inventing fake deterministic steps."],
            ["Accepted", "Draft/proposal is separate from committed reality; graph operations/history support undo."],
            ["Accepted", "Task Type default is used when no estimate is supplied; no actual-effort Q&A or mandatory timer."],
            ["Rejected", "AI black-box auto-decomposition and WBS-level-only stop rules."],
            ["Out of scope", "Project system ownership of general Agent Work Queue or Human Review Queue."],
        ],
        [1800, 7560],
    )

    s.h1("14. MVP acceptance criteria")
    s.bullets([
        "A user can start a Split Session from a node and see parent context plus a visible draft proposal.",
        "The agent supports iterative feedback and proposal versioning; no graph mutation occurs before Commit.",
        "Every proposed leaf is evaluated against all Actionability criteria and includes output/done_when.",
        "Bounded exploration is available for uncertain work.",
        "Commit creates operation history, required/optional contains edges, dependencies, estimates, and an undo reference.",
        "External tool/agent entry points call the same workflow and cannot bypass explicit commit/permission rules.",
        "Post-commit graph health and Elastic Timeline derived values recompute.",
        "User-only Done authority remains intact.",
    ])
    s.p("Reference coverage: v0.1 Sections 7–10, 56–61, and the decomposition-related parts of Sections 58–63 are represented in Sections 1–14. Repetition was consolidated, while workflow states, control meanings, examples, APIs, commit semantics, and non-negotiable rules were retained.")

    s.sources()
    return s.save("03_Collaborative_Task_Decomposition_Spec.docx")


def main() -> None:
    outputs = [build_ui_spec(), build_graph_spec(), build_split_spec()]
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
