"""Build versioned v1.1 addenda from the retained v1.0 DOCX specifications.

The v1.0 documents are opened as templates, never overwritten, and extended with
the decisions approved on 29 Aug 2026. This keeps every earlier requirement and
adds a decision-complete implementation contract.
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from build_specs import (
    ASSETS,
    CALLOUT,
    INK,
    LIGHT_BLUE,
    LIGHT_GRAY,
    MUTED,
    NAVY,
    PAGE_WIDTH_DXA,
    ROOT,
    add_numbering_definition,
    apply_numbering,
    set_paragraph_spacing,
    set_repeat_table_header,
    set_run_font,
    set_table_borders,
    set_table_geometry,
    shade_cell,
)


class Addendum:
    def __init__(self, source: Path, output_name: str, document_label: str):
        self.source = source
        self.output = ROOT / output_name
        self.document_label = document_label
        self.doc = Document(source)
        self.bullet_num_id = add_numbering_definition(self.doc, "bullet")
        self.decimal_num_id = add_numbering_definition(self.doc, "decimal")
        self._update_cover()
        self.doc.add_page_break()

    def _update_cover(self) -> None:
        for paragraph in self.doc.paragraphs[:20]:
            for run in paragraph.runs:
                if "v1.0" in run.text:
                    run.text = run.text.replace("v1.0", "v1.1")
        self.doc.core_properties.version = "1.1"
        self.doc.core_properties.comments = (
            "Versioned refinement. Retains v1.0 content and appends the approved v1.1 contract."
        )

    def h1(self, text: str) -> None:
        self.doc.add_paragraph(text, style="Heading 1")

    def h2(self, text: str) -> None:
        self.doc.add_paragraph(text, style="Heading 2")

    def h3(self, text: str) -> None:
        self.doc.add_paragraph(text, style="Heading 3")

    def p(self, text: str) -> None:
        paragraph = self.doc.add_paragraph()
        set_paragraph_spacing(paragraph)
        run = paragraph.add_run(text)
        set_run_font(run)

    def bullets(self, items: Iterable[str]) -> None:
        for item in items:
            paragraph = self.doc.add_paragraph()
            apply_numbering(paragraph, self.bullet_num_id)
            set_paragraph_spacing(paragraph, after=4, line=1.25)
            set_run_font(paragraph.add_run(item))

    def numbered(self, items: Iterable[str]) -> None:
        for item in items:
            paragraph = self.doc.add_paragraph()
            apply_numbering(paragraph, self.decimal_num_id)
            set_paragraph_spacing(paragraph, after=4, line=1.25)
            set_run_font(paragraph.add_run(item))

    def callout(self, label: str, text: str, fill: str = CALLOUT) -> None:
        table = self.doc.add_table(rows=1, cols=1)
        set_table_geometry(table, [PAGE_WIDTH_DXA])
        set_table_borders(table, color="D9E1E8", size=5)
        cell = table.cell(0, 0)
        shade_cell(cell, fill)
        paragraph = cell.paragraphs[0]
        set_paragraph_spacing(paragraph, after=0)
        set_run_font(paragraph.add_run(f"{label}: "), bold=True, color=NAVY)
        set_run_font(paragraph.add_run(text))
        set_paragraph_spacing(self.doc.add_paragraph(), after=4)

    def code(self, text: str) -> None:
        paragraph = self.doc.add_paragraph()
        set_paragraph_spacing(paragraph, before=4, after=8, line=1.0)
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.right_indent = Inches(0.18)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), LIGHT_GRAY)
        paragraph._p.get_or_add_pPr().append(shading)
        set_run_font(paragraph.add_run(text), name="Consolas", size=9.2, color=INK)

    def table(
        self,
        headers: Sequence[str],
        rows: Sequence[Sequence[str]],
        widths: Sequence[int],
        font_size: float = 9.0,
    ) -> None:
        table = self.doc.add_table(rows=1, cols=len(headers))
        set_table_geometry(table, widths)
        set_table_borders(table)
        header = table.rows[0]
        set_repeat_table_header(header)
        for index, value in enumerate(headers):
            cell = header.cells[index]
            shade_cell(cell, LIGHT_BLUE)
            paragraph = cell.paragraphs[0]
            set_paragraph_spacing(paragraph, after=0, line=1.12)
            set_run_font(paragraph.add_run(value), size=font_size, bold=True, color=NAVY)
        for row_values in rows:
            row = table.add_row()
            for index, value in enumerate(row_values):
                paragraph = row.cells[index].paragraphs[0]
                set_paragraph_spacing(paragraph, after=0, line=1.12)
                set_run_font(paragraph.add_run(str(value)), size=font_size)
        set_table_geometry(table, widths)
        set_paragraph_spacing(self.doc.add_paragraph(), after=4)

    def image(self, path: Path, caption: str, alt_text: str, width: float = 6.5) -> None:
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        shape = paragraph.add_run().add_picture(str(path), width=Inches(width))
        shape._inline.docPr.set("descr", alt_text)
        caption_paragraph = self.doc.add_paragraph(caption, style="Caption")
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def traceability(self, rows: Sequence[Sequence[str]]) -> None:
        self.h2("Source traceability")
        self.table(
            ["Source", "Authority / interpretation", "v1.1 coverage"],
            rows,
            [2050, 3000, 4310],
            font_size=8.6,
        )

    def changelog(self, items: Iterable[str]) -> None:
        self.h2("v1.0 -> v1.1 changelog")
        self.bullets(items)

    def vocabulary(self) -> None:
        self.h2("Normalized cross-document vocabulary")
        self.table(
            ["Concept", "v1.1 values", "Rule"],
            [
                ["NodeKind", "WORK, ARTIFACT, RESOURCE, AGENT", "Describes what kind of entity the node represents."],
                ["WorkType", "UNCLASSIFIED, GOAL, DELIVERABLE, WORK_PACKAGE, ACTION", "Applies only to WORK nodes and replaces ambiguous TASK + WBS combinations."],
                ["Stage", "CAPTURED, PLANNING, READY, EXECUTION, REVIEW, CLOSED", "Describes lifecycle phase."],
                ["Status", "TODO, DOING, BLOCKED, DONE, CANCELLED, SUPERSEDED", "Describes current or terminal runtime truth."],
                ["Schedule", "planned_start, planned_end, deadline", "Planned dates are elastic intent; deadline is an explicit anchor."],
                ["Proposal", "SplitSession + ProposalVersion", "Draft state is never stored as committed graph nodes."],
                ["Operation", "OperationBatch + ordered operations", "All mutations are versioned, atomic, reviewable, and undoable."],
            ],
            [1800, 3100, 4460],
            font_size=8.7,
        )

    def ledgers(
        self,
        accepted: Sequence[str],
        corrected: Sequence[str],
        rejected: Sequence[str],
        deferred: Sequence[str],
        out_of_scope: Sequence[str],
    ) -> None:
        self.h1("Decision ledgers")
        for heading, items in [
            ("Accepted", accepted),
            ("Corrected", corrected),
            ("Rejected", rejected),
            ("Deferred", deferred),
            ("Out of scope", out_of_scope),
        ]:
            self.h2(heading)
            self.bullets(items)

    def save(self) -> Path:
        self.doc.save(self.output)
        return self.output


def build_ui() -> Path:
    spec = Addendum(
        ROOT / "01_Canvas_and_Elastic_Timeline_UI_Spec.docx",
        "01_Canvas_and_Elastic_Timeline_UI_Spec_v1.1.docx",
        "Canvas and Timeline UI",
    )
    spec.h1("v1.1 refinement addendum")
    spec.callout(
        "Version rule",
        "All v1.0 content above remains part of the specification. This addendum resolves later decisions and overrides only where it explicitly says Corrected or Rejected.",
    )
    spec.traceability(
        [
            ["Shared design conversation", "Primary source for user intent, corrections, and non-negotiable behavior.", "All UI sections and decision ledgers."],
            ["Yonc Graph Project System v0.1", "Baseline product, graph, Canvas, Timeline, and Split requirements.", "Contracts, edge cases, acceptance checks."],
            ["UI Spec v1.0", "Retained verbatim as the base specification.", "This addendum narrows unresolved mappings."],
            ["Authoritative Canvas image", "Controls dark shell, graph hierarchy, time axis, inspector, and restrained node styling.", "Final Canvas contract and Figure 4."],
            ["Token-activity image", "Controls calendar heatmap orientation only.", "Capacity Grid week columns, weekday rows, month labels."],
            ["WhatsApp grid image", "Visual reference only: rounded cells, gaps, color regions, dashed selection. Embedded puzzle rules are excluded.", "Capacity Grid cell grammar and Figure 5."],
            ["Current graph_app", "Implementation evidence, not product authority.", "Corrected edge routing, scheduling, state, language, and mobile behavior."],
        ]
    )
    spec.changelog(
        [
            "Canvas horizontal position is now strictly date-derived; vertical position remains view layout.",
            "Timeline is finalized as Forecast + Capacity Grid, with planned-capacity cells rather than completed-activity cells.",
            "Capacity Grid uses week columns, Monday-Sunday rows, month labels, large rounded cells, and split-color overlap.",
            "Main navigation and ordinary field labels remain English; dialogs, confirmations, validation, errors, toasts, and Split conversation use Simplified Chinese.",
            "A final dark Canvas mockup and final Timeline mockup replace unresolved visual alternatives.",
            "Desktop core is the first full milestone; mobile receives a safe compact action view without graph dragging.",
        ]
    )
    spec.vocabulary()

    spec.h1("8. Final Canvas contract")
    spec.image(
        ASSETS / "canvas-ui-final-v1.1.png",
        "Figure 4. Final v1.1 Canvas mockup: date-derived X axis, orthogonal graph, scoped cards, and node inspector.",
        "Dark Yonc Canvas mockup with weekly time columns, orthogonal project graph, progress cards, resource references, deadline signals, minimap, and right-side inspector.",
    )
    spec.h2("8.1 Layout and coordinate ownership")
    spec.bullets(
        [
            "The X coordinate is derived from planned_start, planned_end, or deadline against the visible date scale. It is not an arbitrary saved pixel coordinate.",
            "Unscheduled nodes occupy a dedicated lane before the dated region. The first horizontal drop creates planned_start; the system may derive planned_end from forecast.",
            "The Y coordinate separates project lineage and may be adjusted as view state. Vertical movement never reparents a node.",
            "All committed contains and dependency edges use horizontal/vertical segments with rounded right-angle corners. Curves and diagonals are prohibited.",
            "A deadline node is horizontally locked. Editing the deadline uses an explicit detail flow and Chinese confirmation dialog.",
        ]
    )
    spec.h2("8.2 Card and inspector information")
    spec.table(
        ["Surface", "Always visible", "On selection / detail"],
        [
            ["Collapsed card", "Work type/WBS, title, status perimeter, progress, deadline when present, resource count, warning icon", "No inline long-form editing"],
            ["Inspector", "Selected node identity and current state", "Description, stage/status, effort, span, finish range, deadline, pressure, children, dependencies, resources, warnings, history"],
            ["Neighborhood", "Parents, children, dependencies, and resources remain legible", "Unrelated graph dims but does not disappear"],
            ["Hidden descendants", "Collapsed card reports hidden count", "Expand state persists separately for Canvas and Timeline"],
        ],
        [1700, 3600, 4060],
    )
    spec.h2("8.3 Drag and keyboard behavior")
    spec.bullets(
        [
            "Horizontal drag previews the derived date and validates nearest ancestor deadline plus dependency ordering before apply.",
            "Invalid regions are visibly shaded. A rejected drop restores the prior position and opens a Simplified Chinese explanation.",
            "Keyboard selection, inspector access, expand/collapse, and schedule nudging are available without pointer-only interaction.",
            "The delayed right-edge plus (0.8-1.0 seconds) opens Split Session. On touch/mobile, selection plus an action menu replaces hover.",
            "Reduced-motion mode removes pulsing while retaining static pressure brightness and all warning information.",
        ]
    )

    spec.h1("9. Final Timeline contract")
    spec.image(
        ASSETS / "timeline-grid-cell-style-reference.jpeg",
        "Figure 5. User-supplied cell-style reference. Only rounded cell geometry, spacing, color regions, and dashed selection are adopted; puzzle rules and social-media chrome are excluded.",
        "Reference screenshot with a large rounded-square color grid used only to guide cell shape, gap size, color grouping, and selection outline.",
        width=3.25,
    )
    spec.image(
        ASSETS / "timeline-grid-final-v1.1.png",
        "Figure 6. Final v1.1 Capacity Grid mockup with module pool, calendar heatmap, overlap, today/deadline markers, and forecast summary.",
        "Dark Yonc Capacity Grid with unscheduled modules, horizontal weeks, Monday through Sunday rows, colorful planned-capacity cells, split overlaps, selected date range, and forecast statistics.",
    )
    spec.h2("9.1 Modes, scales, and cell meaning")
    spec.table(
        ["Control", "Meaning"],
        [
            ["Forecast", "Range-based project/module forecast, remaining-work curve, deadline gap, confidence, and pressure explanation."],
            ["Capacity Grid", "Calendar heatmap of planned capacity/intention; it is not a completion or activity heatmap."],
            ["Day / Week / Month / Quarter", "Aggregation/zoom inside Capacity Grid; the canonical day view uses week columns and Monday-Sunday rows."],
            ["Colored cell", "One calendar day containing rough planned allocation for one or more scopes."],
            ["Empty cell", "No planned placement. It makes no claim about productivity or availability."],
        ],
        [2200, 7160],
    )
    spec.h2("9.2 Placement and overlap")
    spec.numbered(
        [
            "Drag an unscheduled L2/L3 module onto a day cell; that cell becomes planned_start.",
            "The system computes an initial natural-calendar range from remaining estimated effort and observed delivery pace.",
            "The user may move or resize the range. Resizing changes planned span, never estimated effort.",
            "Two allocations split a cell 50/50. More than two use compact project-color segments plus a +N indicator.",
            "Overlap raises planning-density warnings but never silently redistributes the user's delivery pace or moves another project.",
        ]
    )
    spec.h2("9.3 Visual states")
    spec.table(
        ["State", "Treatment"],
        [
            ["Selected range", "High-contrast dashed outline around the range, with draggable start/end handles."],
            ["Today", "Single vertical marker and labeled current-day cell; never confused with deadline."],
            ["Deadline", "Amber diamond/flag anchor with explicit date and locked editing flow."],
            ["Invalid date", "Muted red forbidden region plus a Chinese validation explanation before apply."],
            ["Over capacity", "Quiet density warning and details; no moral language or failure framing."],
            ["Low confidence", "Forecast range is withheld or labeled insufficient history rather than invented."],
        ],
        [2200, 7160],
    )

    spec.h1("10. Language, dialogs, and feedback")
    spec.callout(
        "Language rule",
        "Navigation, tabs, field labels, and API terminology remain English. Every modal, confirmation, validation message, error, toast, and Split conversation uses Simplified Chinese.",
    )
    spec.table(
        ["Trigger", "Simplified Chinese dialog / message", "Actions"],
        [
            ["Commit split", "确认提交拆分？提交后将创建正式节点和关系。", "取消 / 提交拆分"],
            ["Discard draft", "放弃当前拆分提案？未提交的内容不会写入项目图。", "继续编辑 / 放弃"],
            ["Deadline edit", "确认修改截止日期？这会重新计算预测和时间压力。", "取消 / 确认修改"],
            ["Invalid drop", "无法放置到此日期：已超出父级截止日期或违反依赖顺序。", "查看原因 / 返回"],
            ["Version conflict", "项目图已发生变化。请刷新上下文并重新确认提案。", "取消 / 刷新"],
            ["Mark done", "确认标记为完成？完成状态只能由你确认。", "取消 / 标记完成"],
            ["Cancel", "请输入取消原因。", "返回 / 确认取消"],
            ["Supersede", "请输入替代原因，并可选择新的替代节点。", "返回 / 确认替代"],
            ["Undo", "已撤销上一批项目图变更。", "关闭"],
            ["Import preview", "导入预览已生成。确认前不会修改本地图。", "取消 / 应用导入"],
            ["Loading failure", "无法载入项目图，请重试。", "重试 / 关闭"],
            ["Insufficient pace", "历史完成数据不足，暂不显示预计完成日期。", "知道了"],
        ],
        [1700, 5300, 2360],
        font_size=8.5,
    )

    spec.h1("11. Responsive, accessibility, and failure states")
    spec.bullets(
        [
            "Desktop is complete. Compact mobile shows Current Scope, next Actions, deadline/forecast, warnings, resources, Split entry, and user-only Done; it does not expose free graph dragging.",
            "Every icon control has an accessible name, visible keyboard focus, and a text equivalent in menus or tooltips.",
            "Color is never the only signal: status, pressure, warning, overlap, and selection also use shape, perimeter, label, or pattern.",
            "Loading uses stable skeleton regions; empty states explain the next valid action; errors preserve unsaved form input when safe.",
            "Large graphs render through scope filtering, semantic zoom, and virtualization; layout computation must not freeze the main interface.",
        ]
    )
    spec.ledgers(
        accepted=[
            "Dark Yonc shell with date-derived Canvas X coordinate and orthogonal edges.",
            "Forecast + Capacity Grid Timeline using calendar heatmap orientation and planned-capacity semantics.",
            "Large rounded grid cells, clear gaps, project color regions, dashed selection, split overlap, and +N overflow.",
            "English main interface with Simplified Chinese dialogs and Split conversation.",
            "Desktop core first with compact action-oriented mobile fallback.",
        ],
        corrected=[
            "Canvas saved pixel X is replaced by date-derived X; only vertical layout remains manual view state.",
            "Current curved implementation edges are replaced by orthogonal routing.",
            "The 14-day row table is replaced by Forecast and calendar Capacity Grid projections.",
        ],
        rejected=[
            "Treating colored grid cells as completed activity.",
            "Importing puzzle adjacency rules, animals, hearts, or social-media controls from the WhatsApp image.",
            "Using glow as a generic attention score or using color alone for state.",
        ],
        deferred=[
            "Full mobile graph editing, light theme, and task-type-specific delivery pace.",
            "User-authored non-contiguous day painting; v1 uses an automatically generated contiguous range that can be moved/resized.",
        ],
        out_of_scope=[
            "Automatic Notion writeback, general agent work queues, human review queues, and autonomous multi-agent scheduling.",
        ],
    )
    spec.h1("12. UI acceptance criteria")
    spec.bullets(
        [
            "Canvas renders scoped graph data with real time guides, date-derived X, orthogonal edges, inspector, minimap, and persistent per-view expansion.",
            "A deadline/dependency-invalid drag cannot mutate the schedule and produces the required Chinese explanation.",
            "Capacity Grid displays horizontal weeks, Monday-Sunday rows, month labels, rounded cells, selection, overlap, today, deadline, and capacity warning states.",
            "Drag-to-start generates a forecast range; move/resize changes dates but not effort.",
            "All dialog/toast/error/Split strings are Simplified Chinese while main navigation remains English.",
            "Desktop and compact-mobile flows are keyboard accessible, readable without color alone, and stable under empty/loading/error states.",
        ]
    )
    return spec.save()


def build_graph() -> Path:
    spec = Addendum(
        ROOT / "02_Graph_Based_Project_Management_System_Spec.docx",
        "02_Graph_Based_Project_Management_System_Spec_v1.1.docx",
        "Graph Project System",
    )
    spec.h1("v1.1 refinement addendum")
    spec.callout(
        "System authority",
        "Local SQLite is the canonical committed project graph. JSON is an export/backup format; Notion is a safe import/projection and receives no automatic writeback in v1.1.",
    )
    spec.traceability(
        [
            ["Shared design conversation", "Primary authority for human control, split, completion, forecast, and scope reductions.", "All normalized rules and ledgers."],
            ["Yonc Graph Project System v0.1", "Baseline graph architecture and non-violation rules.", "Data, operations, forecast, health, resources, agent boundary."],
            ["Graph System Spec v1.0", "Retained verbatim as architectural history and rationale.", "v1.1 addendum resolves implementation contracts."],
            ["Current SQLite graph", "Migration baseline: 551 nodes, 495 edges, four DONE nodes, no planned/deadline dates.", "Migration and zero-loss acceptance."],
            ["Current graph_app API", "Compatibility evidence; `/api/v1` remains temporarily available.", "Corrected schema, operations, pace, proposal, and actor handling."],
        ]
    )
    spec.changelog(
        [
            "SQLite, not repository JSON or Notion, is fixed as committed graph truth.",
            "NodeKind and WorkType are separated; Stage and Status are normalized as independent fields.",
            "contains edges become canonical hierarchy with one incoming committed parent edge.",
            "Versioned SplitSession, ProposalVersion, OperationBatch, ViewState, and ResourceReference records are added.",
            "Forecast, recursive progress, pressure, graph health, error contracts, and migration mappings are made deterministic.",
            "A breaking but compatible `/api/v2` is introduced while `/api/v1` remains available during rollout.",
        ]
    )
    spec.vocabulary()

    spec.h1("13. Canonical persistence and graph integrity")
    spec.bullets(
        [
            "SQLite is authoritative for committed nodes, edges, schedules, resource references, graph version, operation history, and split-session metadata.",
            "Each non-root committed node has at most one incoming contains edge. The API may return derived parent_id, but clients cannot maintain a second hierarchy field.",
            "Every structural write checks source/target existence, self-edge, duplicate relation, one-parent rule, and cycle creation.",
            "Cross-project membership is represented by related/dependency/resource relationships, never a second contains parent.",
            "Normal UI flows never hard-delete committed nodes. CANCELLED, SUPERSEDED, archive metadata, and reversible operation batches preserve history.",
        ]
    )
    spec.h2("13.1 Node contract")
    spec.table(
        ["Field group", "Required v1.1 fields and behavior"],
        [
            ["Identity", "id, title, node_kind, work_type, derived wbs_level, origin, created_at, updated_at"],
            ["Definition", "description, start_cue, inputs, done_when/output definition"],
            ["State", "stage, status, status_reason, closed_from_stage, superseded_by"],
            ["Classification", "tags, task_type, mode, theme; extensible metadata remains JSON"],
            ["Effort", "estimated_effort_minutes, estimate_source (user/AI/task_type_default), optional confidence"],
            ["Time", "planned_start, planned_end, deadline; ISO-8601 dates validated server-side"],
            ["Provenance", "legacy/raw fields retained through migration metadata; Notion block ID remains an external reference"],
        ],
        [2150, 7210],
    )
    spec.h2("13.2 Stage and status rules")
    spec.table(
        ["Action", "Stage result", "Status result", "Constraint"],
        [
            ["Capture", "CAPTURED", "TODO", "Inbox classification may remain UNCLASSIFIED."],
            ["Plan / split", "PLANNING", "TODO", "Proposal remains separate until commit."],
            ["Actionability accepted", "READY", "TODO", "Start cue and done_when are present for Actions."],
            ["Start", "EXECUTION", "DOING", "User action or permitted runtime event."],
            ["Block", "unchanged", "BLOCKED", "Reason/warning may be recorded; stage is not overwritten."],
            ["Submit for review", "REVIEW", "TODO", "Does not imply DONE."],
            ["Done", "CLOSED", "DONE", "User-facing endpoint only; no reason required."],
            ["Cancel / supersede", "CLOSED", "CANCELLED / SUPERSEDED", "Reason required; replacement optional for supersede."],
            ["Reopen / undo", "closed_from_stage", "prior status", "Restores the recorded prior state."],
        ],
        [1700, 1800, 2000, 3860],
        font_size=8.5,
    )

    spec.h1("14. Derived-state contract")
    spec.h2("14.1 Recursive progress")
    spec.code(
        "required_actions = all required descendant WORK/ACTION nodes\n"
        "remaining denominator excludes CANCELLED and SUPERSEDED\n"
        "weight = estimated_effort_minutes after source precedence\n"
        "progress = DONE required-action weight / total required-action weight\n"
        "fallback = completed required Actions / total required Actions"
    )
    spec.p("Calculated progress never sets status. A user may mark a parent DONE while calculated progress is below 100 percent; both facts remain visible.")
    spec.h2("14.2 Observed Delivery Pace and forecast")
    spec.numbered(
        [
            "Use the last eight complete ISO calendar weeks and include zero-output weeks.",
            "Count the latest valid transition into DONE once per node; reopen/undo removes that completion from the active pace sample.",
            "Require at least three completed Actions across at least two distinct weeks and a non-zero median; otherwise return insufficient_history without an invented finish date.",
            "Typical pace is the median weekly completed estimated effort. The 25th/75th weekly pace values drive the slow/fast forecast range.",
            "Remaining effort is required descendant Action effort not in a terminal removed state. Missing values follow user > AI > Task Type default.",
            "Each selected scope is forecast independently in v1. Overlap creates a density warning but does not silently allocate or rebalance global capacity.",
        ]
    )
    spec.h2("14.3 Temporal pressure")
    spec.p("Pressure is derived from remaining effort, forecast finish range, deadline slack, overdue state, and dependency constraints. It is normalized to 0.0-1.0 and accompanied by explainable factors; a fixed days-to-deadline threshold is insufficient.")

    spec.h1("15. Operations, resources, and view state")
    spec.table(
        ["Object", "Required fields / guarantees"],
        [
            ["GraphVersion", "Monotonic integer incremented by each committed OperationBatch."],
            ["OperationBatch", "id, actor_channel, source, graph_version_before/after, ordered operations, inverse operations, created_at, undone_at."],
            ["ResourceReference", "id, node_id, uri, label, resource_type, role, metadata, created_at."],
            ["ViewState", "view, scope_node_id, expanded IDs, selected node, filters, zoom/pan, vertical layout; separated by Canvas/Timeline."],
            ["Schedule", "planned_start/planned_end on node, placement_source, last_user_adjusted_at; grid cells are derived."],
        ],
        [2100, 7260],
    )
    spec.bullets(
        [
            "Batch undo reverses nodes, edges, schedule changes, resources, and state transitions together. Partial undo is rejected.",
            "Position-only Canvas Y updates are view-state writes and do not increment committed graph version.",
            "Resource URIs are validated by scheme and stored as references; promotion to an Artifact node is an explicit graph operation.",
        ]
    )

    spec.h1("16. API v2 contract")
    spec.table(
        ["Interface", "Behavior"],
        [
            ["GET /api/v2/graph", "Scoped committed graph with graph_version, nodes, edges, progress, health, pressure, and forecast summaries."],
            ["POST/PATCH /api/v2/nodes", "Normalized NodeKind/WorkType fields; parent convenience converts to a contains operation."],
            ["POST /api/v2/nodes/{id}/transition", "User-facing state transition; actor is inferred by route/channel and cannot be supplied by the payload."],
            ["POST /api/v2/nodes/{id}/reparent", "Atomic one-parent update with version, cycle, deadline, and dependency validation."],
            ["PUT /api/v2/nodes/{id}/schedule", "Preview/apply planned dates; returns constraint violations without mutation on failure."],
            ["GET /api/v2/timeline", "Date cells, placements, overlaps, capacity density, deadlines, forecasts, and warnings for a range."],
            ["GET/PUT /api/v2/view-state/{view}", "Per-view, per-scope expansion, selection, filters, zoom/pan, and vertical layout."],
            ["/api/v2/split-sessions", "Start, get, message/revise, validate, discard, and atomic commit against graph/proposal version."],
            ["GET /api/v2/operation-batches", "Review history; POST /{id}/undo reverses one complete batch."],
            ["POST /api/v2/import/legacy/preview|apply", "Idempotent preview first; apply only after explicit user confirmation."],
        ],
        [2750, 6610],
        font_size=8.6,
    )
    spec.code('Error response\n{\n  "code": "DEPENDENCY_ORDER_CONFLICT",\n  "message_key": "schedule.dependency_conflict",\n  "params": {"blocking_node_id": "..."}\n}')
    spec.p("The API uses stable English codes and structured parameters. The interface localizes them into Simplified Chinese; raw backend exception text is never used as product copy.")

    spec.h1("17. Migration and compatibility")
    spec.table(
        ["Legacy value", "v1.1 mapping", "Preservation rule"],
        [
            ["kind=TASK", "node_kind=WORK", "Original value retained in migration metadata."],
            ["wbs_level 1/2/3/4", "GOAL / DELIVERABLE / WORK_PACKAGE / ACTION", "Legacy WBS remains derivable and exportable."],
            ["wbs_level missing", "work_type=UNCLASSIFIED", "No inferred actionability."],
            ["TODO", "PLANNING / TODO", "Conservative default unless stronger legacy state evidence exists."],
            ["DOING", "EXECUTION / DOING", "Preserve state event history."],
            ["DONE", "CLOSED / DONE", "Preserve completion event and user authority."],
            ["CANCELLED / SUPERSEDED", "CLOSED / same status", "Reason must remain available; missing legacy reason is flagged, not invented."],
            ["parent_id", "incoming contains edge", "Mismatch aborts migration; API returns derived parent_id for compatibility."],
        ],
        [2300, 3200, 3860],
        font_size=8.6,
    )
    spec.bullets(
        [
            "Create a timestamped SQLite backup before the first Alembic migration.",
            "Validate that all 551 nodes and 495 edges remain represented and all legacy identifiers/tags/links/Notion references survive.",
            "Keep `/api/v1` and the existing UI available as fallback until v2 data, API, UI, and visual tests pass.",
            "Notion import remains repeatable and preview-first. Automatic Notion writeback is disabled.",
        ]
    )

    spec.ledgers(
        accepted=[
            "SQLite canonical truth with versioned graph operations and JSON export/backup.",
            "Separate NodeKind/WorkType and Stage/Status contracts.",
            "One canonical contains parent, atomic Split commits, typed resources, recursive progress, eight-week pace, and explainable pressure.",
            "API v2 with structured error codes and temporary v1 compatibility.",
        ],
        corrected=[
            "parent_id is no longer independently writable from contains edges.",
            "Proposal acceptance no longer creates nodes that remain is_proposed.",
            "Delivery pace uses ISO weeks and valid DONE transitions instead of date-keyed totals.",
            "Client payloads can no longer claim actor=user.",
        ],
        rejected=[
            "Notion-first truth, multiple contains parents, raw JSON mutation, partial Split commit, and binary deadline-only pressure.",
        ],
        deferred=[
            "Automatic Notion writeback, task-type-specific pace, advanced permissions, redo, and full multi-user collaboration.",
        ],
        out_of_scope=[
            "General agent scheduling, agent work queues, human review queues, and autonomous multi-agent orchestration.",
        ],
    )
    spec.h1("18. System acceptance criteria")
    spec.bullets(
        [
            "Migration preserves the current 551 nodes and 495 edges with no lost legacy identifiers, tags, links, WBS data, or Notion references.",
            "One-parent and cycle rules reject invalid writes before graph-version increment.",
            "Split commit is atomic, version-checked, fully undoable, and never leaves committed nodes in proposal state.",
            "Progress is recursive and effort-weighted; DONE remains user-controlled and independent from calculated progress.",
            "Pace, forecast range, pressure, graph health, schedule validation, and low-confidence states are deterministic and tested.",
            "Notion import is idempotent, preview-first, and never writes back automatically.",
        ]
    )
    return spec.save()


def build_split() -> Path:
    spec = Addendum(
        ROOT / "03_Collaborative_Task_Decomposition_Spec.docx",
        "03_Collaborative_Task_Decomposition_Spec_v1.1.docx",
        "Collaborative Task Decomposition",
    )
    spec.h1("v1.1 refinement addendum")
    spec.callout(
        "Commit boundary",
        "Split Session is a versioned proposal service. Conversation and draft trees never mutate the committed graph; only an explicit, version-checked user Commit creates one atomic OperationBatch.",
    )
    spec.traceability(
        [
            ["Shared design conversation", "Primary authority for collaborative decomposition and user approval.", "Session UX, proposal visibility, commit authority."],
            ["Yonc Graph Project System v0.1", "Baseline Actionability Contract, bounded exploration, operations, and external interface.", "Validation, data, examples, acceptance."],
            ["Task Decomposition Spec v1.0", "Retained verbatim as workflow rationale and examples.", "v1.1 adds exact session/version/API/adapter contracts."],
            ["Current proposal API", "Compatibility evidence; it lacks session context/versioning/atomic batch semantics.", "Corrected implementation behavior."],
            ["Language decision", "All Split conversation, confirmations, errors, validation, and toasts are Simplified Chinese.", "Dialog and message contract."],
        ]
    )
    spec.changelog(
        [
            "Split is implemented behind a model-adapter interface, initially connected to the existing DSPy/Gemini pipeline.",
            "Session, message, proposal version, context graph version, validation, commit result, and undo batch are explicit records.",
            "Proposal nodes/edges no longer share committed graph tables or an is_proposed lifecycle flag.",
            "Commit is a single transaction and fails on stale graph/proposal versions or any invalid item.",
            "Split UI and all feedback are finalized in Simplified Chinese; main application labels remain English.",
        ]
    )
    spec.vocabulary()

    spec.h1("15. Split service architecture")
    spec.code(
        "Canvas / Mobile / External Chat / Agent Tool\n"
        "                 |\n"
        "          Split Session API\n"
        "                 |\n"
        "       SplitModelAdapter interface\n"
        "          |              |\n"
        "   DSPy/Gemini v1     future adapter\n"
        "                 |\n"
        " ProposalVersion + validation\n"
        "                 | explicit user Commit\n"
        " OperationBatch -> Graph Core"
    )
    spec.bullets(
        [
            "The adapter receives a structured context snapshot and returns a structured proposal; it cannot call committed graph mutation services.",
            "The first adapter wraps the existing DSPy/Gemini decomposition pipeline. UI and graph operations depend only on the adapter contract.",
            "External agents may start, discuss, or propose through agent-scoped endpoints but cannot use the user-only DONE transition.",
        ]
    )

    spec.h1("16. Exact session and proposal contract")
    spec.table(
        ["Object", "Required fields"],
        [
            ["SplitSession", "id, parent_node_id, state, context_graph_version, context_snapshot, current_proposal_version, created_at, updated_at"],
            ["SplitMessage", "id, session_id, role (user/assistant/system), content, created_at; UI renders the conversation in Simplified Chinese"],
            ["ProposalVersion", "session_id, version, rationale, proposed_nodes, proposed_edges, actionability_results, warnings, created_at"],
            ["Proposed node", "temporary_id, title, node_kind/work_type, description, start_cue, inputs, done_when, estimate, estimate source, tags, required flag"],
            ["Proposed edge", "temporary source/target, relation, required flag, rationale/condition"],
            ["Commit result", "operation_batch_id, graph_version_before/after, created node/edge IDs, warnings, undo reference"],
        ],
        [2100, 7260],
        font_size=8.7,
    )
    spec.h2("16.1 Session states")
    spec.code("OPEN -> PROPOSAL_DRAFT -> PENDING_USER_REVIEW <-> REVISING -> COMMITTED\n                                      \\-> DISCARDED")
    spec.p("APPROVED is not a hidden AI state. The user action that commits the visible current proposal is the approval boundary.")

    spec.h1("17. API and interaction flow")
    spec.table(
        ["Endpoint", "Behavior"],
        [
            ["POST /api/v2/split-sessions", "Start from a committed parent and capture context_graph_version."],
            ["GET /api/v2/split-sessions/{id}", "Return context summary, conversation, current proposal, actionability results, warnings, and state."],
            ["POST /api/v2/split-sessions/{id}/messages", "Store user feedback, call adapter, and create the next ProposalVersion."],
            ["POST /api/v2/split-sessions/{id}/validate", "Run actionability, hierarchy, duplicate, dependency, deadline, and estimate checks without mutation."],
            ["POST /api/v2/split-sessions/{id}/commit", "Require expected_graph_version and proposal_version; create one atomic OperationBatch."],
            ["POST /api/v2/split-sessions/{id}/discard", "Close the session with no committed graph change."],
        ],
        [3200, 6160],
        font_size=8.7,
    )
    spec.numbered(
        [
            "User opens Split from a selected node; the header shows lineage, state, deadline, existing children, and the reason splitting is needed.",
            "The service snapshots graph context/version and obtains a structured proposal through SplitModelAdapter.",
            "The UI shows Chinese conversation beside a live proposal tree with outputs, estimates, required flags, dependencies, validation, and warnings.",
            "Each user revision creates a new immutable ProposalVersion. Earlier versions remain reviewable.",
            "Commit revalidates the current graph and proposal version, creates all nodes/edges in one transaction, increments graph version once, and records one undoable batch.",
        ]
    )

    spec.h1("18. Validation, conflicts, and undo")
    spec.table(
        ["Check", "Failure behavior"],
        [
            ["Actionability", "Every proposed ACTION must have start cue, observable done_when, single intent, bounded effort, and acceptable decision load."],
            ["Hierarchy", "Exactly one proposed contains parent; no orphan, duplicate child, multiple parent, or cycle."],
            ["Dependency", "No self-edge, duplicate, cycle, or impossible ordering."],
            ["Time", "Child schedule/forecast may not silently cross the nearest explicit ancestor deadline."],
            ["Stale graph", "Return GRAPH_VERSION_CONFLICT; do not commit. Chinese UI asks the user to refresh and reconfirm."],
            ["Stale proposal", "Return PROPOSAL_VERSION_CONFLICT; do not commit an older draft."],
            ["Any item failure", "Rollback all proposed nodes, edges, resources, and operations."],
            ["Undo", "Reverse the complete commit batch; restore graph version history and session/commit reference."],
        ],
        [2150, 7210],
        font_size=8.6,
    )

    spec.h1("19. Simplified Chinese Split interface")
    spec.table(
        ["Element", "Required copy / behavior"],
        [
            ["Header", "拆分：{节点标题}；显示当前层级、截止日期与现有子节点。"],
            ["Empty prompt", "告诉我你希望如何拆分，或让我先提出一个版本。"],
            ["Proposal label", "当前提案 v{n} - 尚未写入项目图"],
            ["Continue", "继续讨论"],
            ["Revise", "重新调整"],
            ["Split deeper", "继续拆分"],
            ["Discard", "放弃提案"],
            ["Commit", "提交拆分"],
            ["Validation success", "提案已通过可执行性与项目图检查。"],
            ["Commit success", "拆分已提交，可在 Canvas 和 Timeline 中查看。"],
            ["Conflict", "项目图已发生变化。请刷新上下文并重新确认提案。"],
        ],
        [2300, 7060],
        font_size=8.7,
    )

    spec.h1("20. Model adapter contract")
    spec.code(
        "class SplitModelAdapter:\n"
        "    propose(context, user_message, previous_proposal) -> ProposalDraft\n"
        "    revise(context, conversation, previous_proposal) -> ProposalDraft\n"
        "\n"
        "ProposalDraft contains structured nodes, edges, rationale, actionability results, and warnings."
    )
    spec.bullets(
        [
            "Adapter output is schema-validated before persistence as a ProposalVersion.",
            "Model failure preserves conversation and the last good proposal; Chinese UI offers retry without duplicate versions.",
            "The adapter receives only the scoped context needed for decomposition and never database/session objects.",
            "A deterministic test adapter supports API and UI tests without live model calls.",
        ]
    )

    spec.ledgers(
        accepted=[
            "Versioned collaborative Split Session with a visible current proposal and explicit user Commit.",
            "Model-agnostic adapter with the existing DSPy/Gemini pipeline as the first implementation.",
            "Separate proposal storage, graph/proposal optimistic concurrency, atomic OperationBatch, and complete undo.",
            "Simplified Chinese Split conversation and feedback.",
        ],
        corrected=[
            "The existing generic Proposal accept/reject queue is not sufficient for Split Session semantics.",
            "Accepted drafts become ordinary committed nodes; is_proposed never survives commit.",
            "Actor identity is inferred by endpoint/channel rather than trusted from client JSON.",
        ],
        rejected=[
            "Black-box automatic decomposition, silent mutation, partial proposal acceptance, and AI-controlled DONE.",
        ],
        deferred=[
            "Multiple model routing, advanced permissions, simultaneous multi-user editing, and richer proposal diff visualization.",
        ],
        out_of_scope=[
            "General agent work queues, human review queues, workforce scheduling, and autonomous multi-agent orchestration.",
        ],
    )
    spec.h1("21. Split acceptance criteria")
    spec.bullets(
        [
            "Starting Split captures the committed parent context and graph version without mutating graph truth.",
            "Each feedback round creates a visible immutable proposal version through the adapter contract.",
            "Every proposed Action includes start cue, done_when, estimate source, required flag, and Actionability results.",
            "Commit fails cleanly on stale graph/proposal version or any validation error and creates no partial nodes/edges.",
            "Successful commit creates one graph-version increment, one OperationBatch, committed nodes without proposal styling, and one undo reference.",
            "All Split conversation, validation, confirmation, success, and failure feedback is Simplified Chinese.",
            "Agents may propose but cannot bypass user Commit or user-only DONE authority.",
        ]
    )
    return spec.save()


def main() -> None:
    for output in (build_ui(), build_graph(), build_split()):
        print(output)


if __name__ == "__main__":
    main()
